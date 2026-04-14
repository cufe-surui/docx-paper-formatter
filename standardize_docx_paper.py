from __future__ import annotations

import copy
import json
import re
from pathlib import Path
from typing import Iterable, Iterator, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


CONFIG_PATH = Path(__file__).with_name("formatter_config.jsonc")
LEGACY_CONFIG_PATH = Path(__file__).with_name("formatter_config.json")

BASE_DEFAULT_CONFIG = {
    "general": {
        "suffix": "_标准论文格式",
        "default_output_subdir": "output/doc",
        "overwrite": False,
        "add_header_footer": True,
    },
    "fonts": {
        "western": "Times New Roman",
        "chinese": "宋体",
    },
    "sizes_pt": {
        "title": 15,
        "abstract_heading": 16,
        "level1_heading": 14,
        "body": 12,
        "small_five": 10.5,
    },
    "indents_cm": {
        "two_char": 0.74,
    },
    "page": {
        "width_cm": 21,
        "height_cm": 29.7,
        "top_margin_cm": 2.54,
        "bottom_margin_cm": 2.54,
        "left_margin_cm": 2.54,
        "right_margin_cm": 2.54,
        "header_distance_cm": 1.5,
        "footer_distance_cm": 1.5,
    },
    "spacing_pt": {
        "title_after": 12,
        "abstract_heading_after": 12,
        "level1_before": 12,
        "level1_after": 12,
        "reference_heading_before": 24,
        "reference_heading_after": 12,
        "caption_before": 12,
        "caption_after": 12,
        "equation_before": 12,
        "equation_after": 12,
    },
    "labels": {
        "abstract_cn": "摘　要",
        "abstract_en": "Abstract",
        "keywords_cn": "关键词：",
        "keywords_en": "Keywords:",
        "references": "参考文献",
    },
}

DEFAULT_SUFFIX = BASE_DEFAULT_CONFIG["general"]["suffix"]
DEFAULT_OUTPUT_SUBDIR = BASE_DEFAULT_CONFIG["general"]["default_output_subdir"]
DEFAULT_WESTERN_FONT = BASE_DEFAULT_CONFIG["fonts"]["western"]
DEFAULT_CHINESE_FONT = BASE_DEFAULT_CONFIG["fonts"]["chinese"]
TITLE_SIZE_PT = BASE_DEFAULT_CONFIG["sizes_pt"]["title"]
ABSTRACT_HEADING_SIZE_PT = BASE_DEFAULT_CONFIG["sizes_pt"]["abstract_heading"]
LEVEL1_HEADING_SIZE_PT = BASE_DEFAULT_CONFIG["sizes_pt"]["level1_heading"]
BODY_SIZE_PT = BASE_DEFAULT_CONFIG["sizes_pt"]["body"]
SMALL_FIVE_PT = BASE_DEFAULT_CONFIG["sizes_pt"]["small_five"]
TWO_CHAR_INDENT_CM = BASE_DEFAULT_CONFIG["indents_cm"]["two_char"]
PAGE_WIDTH_CM = BASE_DEFAULT_CONFIG["page"]["width_cm"]
PAGE_HEIGHT_CM = BASE_DEFAULT_CONFIG["page"]["height_cm"]
TOP_MARGIN_CM = BASE_DEFAULT_CONFIG["page"]["top_margin_cm"]
BOTTOM_MARGIN_CM = BASE_DEFAULT_CONFIG["page"]["bottom_margin_cm"]
LEFT_MARGIN_CM = BASE_DEFAULT_CONFIG["page"]["left_margin_cm"]
RIGHT_MARGIN_CM = BASE_DEFAULT_CONFIG["page"]["right_margin_cm"]
HEADER_DISTANCE_CM = BASE_DEFAULT_CONFIG["page"]["header_distance_cm"]
FOOTER_DISTANCE_CM = BASE_DEFAULT_CONFIG["page"]["footer_distance_cm"]
TITLE_SPACE_AFTER_PT = BASE_DEFAULT_CONFIG["spacing_pt"]["title_after"]
ABSTRACT_HEADING_SPACE_AFTER_PT = BASE_DEFAULT_CONFIG["spacing_pt"]["abstract_heading_after"]
LEVEL1_SPACE_BEFORE_PT = BASE_DEFAULT_CONFIG["spacing_pt"]["level1_before"]
LEVEL1_SPACE_AFTER_PT = BASE_DEFAULT_CONFIG["spacing_pt"]["level1_after"]
REFERENCE_HEADING_SPACE_BEFORE_PT = BASE_DEFAULT_CONFIG["spacing_pt"]["reference_heading_before"]
REFERENCE_HEADING_SPACE_AFTER_PT = BASE_DEFAULT_CONFIG["spacing_pt"]["reference_heading_after"]
CAPTION_SPACE_BEFORE_PT = BASE_DEFAULT_CONFIG["spacing_pt"]["caption_before"]
CAPTION_SPACE_AFTER_PT = BASE_DEFAULT_CONFIG["spacing_pt"]["caption_after"]
EQUATION_SPACE_BEFORE_PT = BASE_DEFAULT_CONFIG["spacing_pt"]["equation_before"]
EQUATION_SPACE_AFTER_PT = BASE_DEFAULT_CONFIG["spacing_pt"]["equation_after"]
ABSTRACT_CN_LABEL = BASE_DEFAULT_CONFIG["labels"]["abstract_cn"]
ABSTRACT_EN_LABEL = BASE_DEFAULT_CONFIG["labels"]["abstract_en"]
KEYWORDS_CN_LABEL = BASE_DEFAULT_CONFIG["labels"]["keywords_cn"]
KEYWORDS_EN_LABEL = BASE_DEFAULT_CONFIG["labels"]["keywords_en"]
REFERENCES_LABEL = BASE_DEFAULT_CONFIG["labels"]["references"]

FIRST_LEVEL_HEADING_RE = re.compile(r"^[一二三四五六七八九十百]+、")
SECOND_LEVEL_HEADING_RE = re.compile(r"^（+[一二三四五六七八九十百]+）")
THIRD_LEVEL_HEADING_RE = re.compile(r"^\d+[．.]")
FOURTH_LEVEL_HEADING_RE = re.compile(r"^（\d+）")
FIFTH_LEVEL_HEADING_RE = re.compile(r"^[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮]")
TABLE_CAPTION_RE = re.compile(r"^表\s*\d+")
FIGURE_CAPTION_RE = re.compile(r"^图\s*\d+")
SOURCE_NOTE_RE = re.compile(r"^(资料来源|数据来源|注：)")


def deep_merge(base: dict, overrides: dict) -> dict:
    merged = copy.deepcopy(base)
    for key, value in overrides.items():
        if key in merged and isinstance(merged[key], dict) and isinstance(value, dict):
            merged[key] = deep_merge(merged[key], value)
        else:
            merged[key] = value
    return merged


def default_config() -> dict:
    return copy.deepcopy(BASE_DEFAULT_CONFIG)


def strip_jsonc_comments(text: str) -> str:
    result: list[str] = []
    in_string = False
    escape = False
    in_line_comment = False
    in_block_comment = False
    index = 0

    while index < len(text):
        char = text[index]
        next_char = text[index + 1] if index + 1 < len(text) else ""

        if in_line_comment:
            if char == "\n":
                in_line_comment = False
                result.append(char)
            index += 1
            continue

        if in_block_comment:
            if char == "*" and next_char == "/":
                in_block_comment = False
                index += 2
            else:
                index += 1
            continue

        if in_string:
            result.append(char)
            if escape:
                escape = False
            elif char == "\\":
                escape = True
            elif char == '"':
                in_string = False
            index += 1
            continue

        if char == "/" and next_char == "/":
            in_line_comment = True
            index += 2
            continue

        if char == "/" and next_char == "*":
            in_block_comment = True
            index += 2
            continue

        result.append(char)
        if char == '"':
            in_string = True
        index += 1

    return "".join(result)


def parse_config_text(text: str) -> dict:
    stripped = strip_jsonc_comments(text).strip()
    if not stripped:
        return {}
    return json.loads(stripped)


def render_config_text(config: dict) -> str:
    general = config["general"]
    fonts = config["fonts"]
    sizes = config["sizes_pt"]
    indents = config["indents_cm"]
    page = config["page"]
    spacing = config["spacing_pt"]
    labels = config["labels"]

    return f"""{{\n  // general: GUI 运行选项\n  "general": {{\n    // suffix: 生成后的文档文件名后缀\n    "suffix": {json.dumps(general["suffix"], ensure_ascii=False)},\n    // default_output_subdir: 未手动指定导出目录时，默认输出到输入文档所在位置下的哪个子目录\n    "default_output_subdir": {json.dumps(general["default_output_subdir"], ensure_ascii=False)},\n    // overwrite: 输出文件已存在时，是否允许直接覆盖\n    "overwrite": {json.dumps(general["overwrite"])},\n    // add_header_footer: 是否自动写入页眉标题和页脚页码\n    "add_header_footer": {json.dumps(general["add_header_footer"])}\n  }},\n\n  // fonts: 中英文字体设置\n  "fonts": {{\n    // western: 英文、数字和西文字母默认字体\n    "western": {json.dumps(fonts["western"], ensure_ascii=False)},\n    // chinese: 中文默认字体\n    "chinese": {json.dumps(fonts["chinese"], ensure_ascii=False)}\n  }},\n\n  // sizes_pt: 各类文本的字号，单位为磅 pt\n  "sizes_pt": {{\n    // title: 论文题目字号\n    "title": {json.dumps(sizes["title"])},\n    // abstract_heading: “摘要/Abstract”标题字号\n    "abstract_heading": {json.dumps(sizes["abstract_heading"])},\n    // level1_heading: 一级标题字号\n    "level1_heading": {json.dumps(sizes["level1_heading"])},\n    // body: 正文、摘要正文、参考文献正文等常规内容字号\n    "body": {json.dumps(sizes["body"])},\n    // small_five: 图表说明、页眉页脚等较小文字字号\n    "small_five": {json.dumps(sizes["small_five"])}\n  }},\n\n  // indents_cm: 缩进设置，单位为厘米 cm\n  "indents_cm": {{\n    // two_char: 中文常用“两字符缩进”的近似宽度\n    "two_char": {json.dumps(indents["two_char"])}\n  }},\n\n  // page: 页面尺寸和页边距设置，单位为厘米 cm\n  "page": {{\n    // width_cm: 页面宽度，A4 通常为 21\n    "width_cm": {json.dumps(page["width_cm"])},\n    // height_cm: 页面高度，A4 通常为 29.7\n    "height_cm": {json.dumps(page["height_cm"])},\n    // top_margin_cm: 上页边距\n    "top_margin_cm": {json.dumps(page["top_margin_cm"])},\n    // bottom_margin_cm: 下页边距\n    "bottom_margin_cm": {json.dumps(page["bottom_margin_cm"])},\n    // left_margin_cm: 左页边距\n    "left_margin_cm": {json.dumps(page["left_margin_cm"])},\n    // right_margin_cm: 右页边距\n    "right_margin_cm": {json.dumps(page["right_margin_cm"])},\n    // header_distance_cm: 页眉到页面顶端的距离\n    "header_distance_cm": {json.dumps(page["header_distance_cm"])},\n    // footer_distance_cm: 页脚到页面底端的距离\n    "footer_distance_cm": {json.dumps(page["footer_distance_cm"])}\n  }},\n\n  // spacing_pt: 各类段落前后间距，单位为磅 pt\n  "spacing_pt": {{\n    // title_after: 论文题目后的段后间距\n    "title_after": {json.dumps(spacing["title_after"])},\n    // abstract_heading_after: 摘要标题后的段后间距\n    "abstract_heading_after": {json.dumps(spacing["abstract_heading_after"])},\n    // level1_before: 一级标题前的段前间距\n    "level1_before": {json.dumps(spacing["level1_before"])},\n    // level1_after: 一级标题后的段后间距\n    "level1_after": {json.dumps(spacing["level1_after"])},\n    // reference_heading_before: “参考文献”标题前的段前间距\n    "reference_heading_before": {json.dumps(spacing["reference_heading_before"])},\n    // reference_heading_after: “参考文献”标题后的段后间距\n    "reference_heading_after": {json.dumps(spacing["reference_heading_after"])},\n    // caption_before: 图表标题前的段前间距\n    "caption_before": {json.dumps(spacing["caption_before"])},\n    // caption_after: 图表标题后的段后间距\n    "caption_after": {json.dumps(spacing["caption_after"])},\n    // equation_before: 独立公式前的段前间距\n    "equation_before": {json.dumps(spacing["equation_before"])},\n    // equation_after: 独立公式后的段后间距\n    "equation_after": {json.dumps(spacing["equation_after"])}\n  }},\n\n  // labels: 自动写入的固定标签文本\n  "labels": {{\n    // abstract_cn: 中文摘要标题文本\n    "abstract_cn": {json.dumps(labels["abstract_cn"], ensure_ascii=False)},\n    // abstract_en: 英文摘要标题文本\n    "abstract_en": {json.dumps(labels["abstract_en"], ensure_ascii=False)},\n    // keywords_cn: 中文关键词标签文本\n    "keywords_cn": {json.dumps(labels["keywords_cn"], ensure_ascii=False)},\n    // keywords_en: 英文关键词标签文本\n    "keywords_en": {json.dumps(labels["keywords_en"], ensure_ascii=False)},\n    // references: 参考文献标题文本\n    "references": {json.dumps(labels["references"], ensure_ascii=False)}\n  }}\n}}\n"""


def ensure_config_file(config_path: Path) -> None:
    if config_path.exists():
        return
    config_path.parent.mkdir(parents=True, exist_ok=True)
    if config_path == CONFIG_PATH and LEGACY_CONFIG_PATH.exists():
        legacy_data = json.loads(LEGACY_CONFIG_PATH.read_text(encoding="utf-8"))
        config_path.write_text(render_config_text(deep_merge(default_config(), legacy_data)), encoding="utf-8")
        return
    config_path.write_text(render_config_text(default_config()), encoding="utf-8")


def load_config(config_path: Path | None = None) -> tuple[Path, dict]:
    resolved_path = (config_path or CONFIG_PATH).expanduser().resolve()
    ensure_config_file(resolved_path)
    raw_data = parse_config_text(resolved_path.read_text(encoding="utf-8"))
    config = deep_merge(default_config(), raw_data)
    return resolved_path, config


def load_config_text(config_path: Path | None = None) -> tuple[Path, str, dict]:
    resolved_path, config = load_config(config_path)
    return resolved_path, resolved_path.read_text(encoding="utf-8"), config


def save_config(config: dict, config_path: Path | None = None) -> Path:
    resolved_path = (config_path or CONFIG_PATH).expanduser().resolve()
    resolved_path.parent.mkdir(parents=True, exist_ok=True)
    resolved_path.write_text(render_config_text(deep_merge(default_config(), config)), encoding="utf-8")
    return resolved_path


def save_config_text(text: str, config_path: Path | None = None) -> tuple[Path, dict]:
    parsed = parse_config_text(text)
    merged = deep_merge(default_config(), parsed)
    resolved_path = save_config(merged, config_path)
    return resolved_path, merged


def default_config_text() -> str:
    return render_config_text(default_config())


def normalize_config_text(text: str) -> str:
    return render_config_text(deep_merge(default_config(), parse_config_text(text)))


def apply_runtime_config(config: dict) -> None:
    global DEFAULT_SUFFIX, DEFAULT_OUTPUT_SUBDIR
    global DEFAULT_WESTERN_FONT, DEFAULT_CHINESE_FONT
    global TITLE_SIZE_PT, ABSTRACT_HEADING_SIZE_PT, LEVEL1_HEADING_SIZE_PT, BODY_SIZE_PT, SMALL_FIVE_PT
    global TWO_CHAR_INDENT_CM
    global PAGE_WIDTH_CM, PAGE_HEIGHT_CM, TOP_MARGIN_CM, BOTTOM_MARGIN_CM, LEFT_MARGIN_CM, RIGHT_MARGIN_CM
    global HEADER_DISTANCE_CM, FOOTER_DISTANCE_CM
    global TITLE_SPACE_AFTER_PT, ABSTRACT_HEADING_SPACE_AFTER_PT, LEVEL1_SPACE_BEFORE_PT, LEVEL1_SPACE_AFTER_PT
    global REFERENCE_HEADING_SPACE_BEFORE_PT, REFERENCE_HEADING_SPACE_AFTER_PT
    global CAPTION_SPACE_BEFORE_PT, CAPTION_SPACE_AFTER_PT, EQUATION_SPACE_BEFORE_PT, EQUATION_SPACE_AFTER_PT
    global ABSTRACT_CN_LABEL, ABSTRACT_EN_LABEL, KEYWORDS_CN_LABEL, KEYWORDS_EN_LABEL, REFERENCES_LABEL

    DEFAULT_SUFFIX = config["general"]["suffix"]
    DEFAULT_OUTPUT_SUBDIR = config["general"]["default_output_subdir"]
    DEFAULT_WESTERN_FONT = config["fonts"]["western"]
    DEFAULT_CHINESE_FONT = config["fonts"]["chinese"]
    TITLE_SIZE_PT = config["sizes_pt"]["title"]
    ABSTRACT_HEADING_SIZE_PT = config["sizes_pt"]["abstract_heading"]
    LEVEL1_HEADING_SIZE_PT = config["sizes_pt"]["level1_heading"]
    BODY_SIZE_PT = config["sizes_pt"]["body"]
    SMALL_FIVE_PT = config["sizes_pt"]["small_five"]
    TWO_CHAR_INDENT_CM = config["indents_cm"]["two_char"]
    PAGE_WIDTH_CM = config["page"]["width_cm"]
    PAGE_HEIGHT_CM = config["page"]["height_cm"]
    TOP_MARGIN_CM = config["page"]["top_margin_cm"]
    BOTTOM_MARGIN_CM = config["page"]["bottom_margin_cm"]
    LEFT_MARGIN_CM = config["page"]["left_margin_cm"]
    RIGHT_MARGIN_CM = config["page"]["right_margin_cm"]
    HEADER_DISTANCE_CM = config["page"]["header_distance_cm"]
    FOOTER_DISTANCE_CM = config["page"]["footer_distance_cm"]
    TITLE_SPACE_AFTER_PT = config["spacing_pt"]["title_after"]
    ABSTRACT_HEADING_SPACE_AFTER_PT = config["spacing_pt"]["abstract_heading_after"]
    LEVEL1_SPACE_BEFORE_PT = config["spacing_pt"]["level1_before"]
    LEVEL1_SPACE_AFTER_PT = config["spacing_pt"]["level1_after"]
    REFERENCE_HEADING_SPACE_BEFORE_PT = config["spacing_pt"]["reference_heading_before"]
    REFERENCE_HEADING_SPACE_AFTER_PT = config["spacing_pt"]["reference_heading_after"]
    CAPTION_SPACE_BEFORE_PT = config["spacing_pt"]["caption_before"]
    CAPTION_SPACE_AFTER_PT = config["spacing_pt"]["caption_after"]
    EQUATION_SPACE_BEFORE_PT = config["spacing_pt"]["equation_before"]
    EQUATION_SPACE_AFTER_PT = config["spacing_pt"]["equation_after"]
    ABSTRACT_CN_LABEL = config["labels"]["abstract_cn"]
    ABSTRACT_EN_LABEL = config["labels"]["abstract_en"]
    KEYWORDS_CN_LABEL = config["labels"]["keywords_cn"]
    KEYWORDS_EN_LABEL = config["labels"]["keywords_en"]
    REFERENCES_LABEL = config["labels"]["references"]


def iter_input_files(paths: Sequence[str | Path], suffix: str) -> Iterator[Path]:
    seen: set[Path] = set()
    for raw_path in paths:
        path = Path(raw_path).expanduser().resolve()
        if path.is_file():
            candidates = [path]
        else:
            continue

        for candidate in candidates:
            if candidate in seen:
                continue
            seen.add(candidate)
            if is_source_docx(candidate, suffix):
                yield candidate


def is_source_docx(path: Path, suffix: str) -> bool:
    return (
        path.is_file()
        and path.suffix.lower() == ".docx"
        and not path.name.startswith("~$")
        and not path.name.startswith(".~")
        and not path.stem.endswith(suffix)
    )


def resolve_output_dir(input_files: Sequence[Path], output_dir: Path | None) -> Path:
    if output_dir is not None:
        return output_dir.expanduser().resolve()

    unique_parents = {path.parent for path in input_files}
    if not unique_parents:
        raise ValueError("未提供有效的输入文档路径。")
    base_dir = next(iter(unique_parents)) if len(unique_parents) == 1 else next(iter(unique_parents))
    return (base_dir / Path(DEFAULT_OUTPUT_SUBDIR)).resolve()


def non_empty_paragraphs(document: Document) -> list:
    return [paragraph for paragraph in document.paragraphs if paragraph_has_visible_content(paragraph)]


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()


def normalized_compact(text: str) -> str:
    return normalize_text(text).replace(" ", "")


def looks_like_english_title(text: str) -> bool:
    normalized = normalize_text(text)
    if not normalized:
        return False
    english_letters = sum(1 for char in normalized if char.isascii() and char.isalpha())
    chinese_chars = sum(1 for char in normalized if "\u4e00" <= char <= "\u9fff")
    return english_letters >= 12 and english_letters > chinese_chars


def clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def clear_paragraph_content(paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def set_run_font(
    run,
    east_asia_font: str,
    western_font: str,
    size_pt: float,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = western_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), western_font)
    r_fonts.set(qn("w:hAnsi"), western_font)
    r_fonts.set(qn("w:cs"), western_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def rewrite_paragraph(
    paragraph,
    text: str,
    east_asia_font: str,
    western_font: str,
    size_pt: float,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    clear_paragraph_runs(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia_font, western_font, size_pt, bold=bold, italic=italic)


def style_existing_runs(
    paragraph,
    east_asia_font: str,
    western_font: str,
    size_pt: float,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    if not paragraph.runs and normalize_text(paragraph.text):
        rewrite_paragraph(
            paragraph,
            normalize_text(paragraph.text),
            east_asia_font,
            western_font,
            size_pt,
            bold=bold,
            italic=italic,
        )
        return

    for run in paragraph.runs:
        set_run_font(run, east_asia_font, western_font, size_pt, bold=bold, italic=italic)


def set_paragraph_layout(
    paragraph,
    *,
    alignment,
    first_line_indent_cm: float = 0,
    left_indent_cm: float = 0,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
    line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
) -> None:
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(first_line_indent_cm)
    fmt.left_indent = Cm(left_indent_cm)
    fmt.right_indent = Cm(0)
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    fmt.line_spacing_rule = line_spacing_rule


def configure_page(document: Document) -> None:
    for section in document.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(TOP_MARGIN_CM)
        section.bottom_margin = Cm(BOTTOM_MARGIN_CM)
        section.left_margin = Cm(LEFT_MARGIN_CM)
        section.right_margin = Cm(RIGHT_MARGIN_CM)
        section.header_distance = Cm(HEADER_DISTANCE_CM)
        section.footer_distance = Cm(FOOTER_DISTANCE_CM)


def configure_normal_style(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = DEFAULT_WESTERN_FONT
    normal_style.font.size = Pt(BODY_SIZE_PT)
    style_element = normal_style.element
    r_pr = style_element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), DEFAULT_WESTERN_FONT)
    r_fonts.set(qn("w:hAnsi"), DEFAULT_WESTERN_FONT)
    r_fonts.set(qn("w:cs"), DEFAULT_WESTERN_FONT)
    r_fonts.set(qn("w:eastAsia"), DEFAULT_CHINESE_FONT)


def paragraph_child_tags(paragraph) -> list[str]:
    return [child.tag.split("}")[-1] for child in paragraph._element]


def paragraph_has_math(paragraph) -> bool:
    return any(tag in {"oMath", "oMathPara"} for tag in paragraph_child_tags(paragraph))


def paragraph_has_visible_content(paragraph) -> bool:
    if normalize_text(paragraph.text):
        return True
    return any(tag in {"oMath", "oMathPara", "drawing", "object", "pict"} for tag in paragraph_child_tags(paragraph))


def find_paragraph_index(paragraphs: Sequence, predicate) -> int | None:
    for index, paragraph in enumerate(paragraphs):
        if predicate(paragraph):
            return index
    return None


def next_index(candidates: Iterable[int | None]) -> int | None:
    values = [value for value in candidates if value is not None]
    return min(values) if values else None


def is_abstract_heading(text: str) -> bool:
    return normalized_compact(text) == "摘要"


def is_english_abstract_heading(text: str) -> bool:
    return normalized_compact(text).lower() == "abstract"


def is_keywords_line(text: str) -> bool:
    compact = normalized_compact(text)
    return compact.startswith("关键词") or compact.startswith("关键字")


def is_english_keywords_line(text: str) -> bool:
    compact = normalized_compact(text).lower()
    return compact.startswith("keywords") or compact.startswith("keyword")


def is_reference_heading(text: str) -> bool:
    return normalized_compact(text) == "参考文献"


def split_keywords(text: str) -> list[str]:
    if "：" in text:
        _, keyword_body = text.split("：", 1)
    elif ":" in text:
        _, keyword_body = text.split(":", 1)
    else:
        keyword_body = text
    candidates = re.split(r"[；;，,、]\s*", keyword_body)
    return [item.strip() for item in candidates if item.strip()]


def apply_title_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=TITLE_SPACE_AFTER_PT,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    style_existing_runs(paragraph, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, TITLE_SIZE_PT, bold=True)


def apply_english_title_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=TITLE_SPACE_AFTER_PT,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    style_existing_runs(paragraph, DEFAULT_WESTERN_FONT, DEFAULT_WESTERN_FONT, TITLE_SIZE_PT, bold=True)


def apply_abstract_heading_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=ABSTRACT_HEADING_SPACE_AFTER_PT,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    rewrite_paragraph(paragraph, ABSTRACT_CN_LABEL, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, ABSTRACT_HEADING_SIZE_PT, bold=True)


def apply_english_abstract_heading_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=ABSTRACT_HEADING_SPACE_AFTER_PT,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    rewrite_paragraph(paragraph, ABSTRACT_EN_LABEL, DEFAULT_WESTERN_FONT, DEFAULT_WESTERN_FONT, ABSTRACT_HEADING_SIZE_PT, bold=True)


def apply_abstract_body_format(paragraph, *, english: bool) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_cm=0,
        left_indent_cm=0,
        space_after_pt=0,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
    )
    font_name = DEFAULT_WESTERN_FONT if english else DEFAULT_CHINESE_FONT
    style_existing_runs(paragraph, font_name, DEFAULT_WESTERN_FONT, BODY_SIZE_PT)


def apply_keywords_format(paragraph, *, english: bool) -> None:
    items = split_keywords(paragraph.text)
    label = KEYWORDS_EN_LABEL if english else KEYWORDS_CN_LABEL
    body = "  ".join(items) if english else "\u3000\u3000".join(items)

    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        left_indent_cm=TWO_CHAR_INDENT_CM,
        first_line_indent_cm=0,
        space_after_pt=0,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
    )
    clear_paragraph_runs(paragraph)

    font_name = DEFAULT_WESTERN_FONT if english else DEFAULT_CHINESE_FONT
    label_run = paragraph.add_run(label)
    set_run_font(label_run, font_name, DEFAULT_WESTERN_FONT, BODY_SIZE_PT, bold=True)

    body_run = paragraph.add_run(body)
    set_run_font(body_run, font_name, DEFAULT_WESTERN_FONT, BODY_SIZE_PT)


def apply_level1_heading_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=LEVEL1_SPACE_BEFORE_PT,
        space_after_pt=LEVEL1_SPACE_AFTER_PT,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    style_existing_runs(paragraph, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, LEVEL1_HEADING_SIZE_PT, bold=True)


def apply_level2_heading_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        left_indent_cm=TWO_CHAR_INDENT_CM,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    style_existing_runs(paragraph, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, BODY_SIZE_PT, bold=True)


def apply_level3_heading_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        left_indent_cm=TWO_CHAR_INDENT_CM,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    style_existing_runs(paragraph, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, BODY_SIZE_PT)


def apply_run_in_heading_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        left_indent_cm=TWO_CHAR_INDENT_CM,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
    )
    style_existing_runs(paragraph, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, BODY_SIZE_PT)


def apply_reference_heading_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=REFERENCE_HEADING_SPACE_BEFORE_PT,
        space_after_pt=REFERENCE_HEADING_SPACE_AFTER_PT,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    rewrite_paragraph(paragraph, REFERENCES_LABEL, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, LEVEL1_HEADING_SIZE_PT, bold=True)


def apply_reference_item_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_cm=0,
        left_indent_cm=0,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
    )
    style_existing_runs(paragraph, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, BODY_SIZE_PT)


def apply_body_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_cm=TWO_CHAR_INDENT_CM,
        left_indent_cm=0,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
    )
    style_existing_runs(paragraph, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, BODY_SIZE_PT)


def apply_caption_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before_pt=CAPTION_SPACE_BEFORE_PT,
        space_after_pt=CAPTION_SPACE_AFTER_PT,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    style_existing_runs(paragraph, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, SMALL_FIVE_PT)


def apply_source_note_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    style_existing_runs(paragraph, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, SMALL_FIVE_PT)


def apply_equation_block_format(paragraph) -> None:
    set_paragraph_layout(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_cm=0,
        left_indent_cm=0,
        space_before_pt=EQUATION_SPACE_BEFORE_PT,
        space_after_pt=EQUATION_SPACE_AFTER_PT,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
    )
    style_existing_runs(paragraph, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, BODY_SIZE_PT)


def format_table_cells(document: Document) -> None:
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not normalize_text(paragraph.text):
                        continue
                    set_paragraph_layout(
                        paragraph,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        first_line_indent_cm=0,
                        left_indent_cm=0,
                        line_spacing_rule=WD_LINE_SPACING.SINGLE,
                    )
                    rewrite_paragraph(paragraph, normalize_text(paragraph.text), DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, SMALL_FIVE_PT)


def set_header_footer(document: Document, title_text: str, enabled: bool) -> None:
    if not enabled or not title_text:
        return

    for section in document.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        header_paragraph = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        clear_paragraph_content(header_paragraph)
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_paragraph.add_run(title_text)
        set_run_font(header_run, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, SMALL_FIVE_PT)

        footer_paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        clear_paragraph_content(footer_paragraph)
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        append_page_field(footer_paragraph)


def append_page_field(paragraph) -> None:
    run = paragraph.add_run()
    set_run_font(run, DEFAULT_CHINESE_FONT, DEFAULT_WESTERN_FONT, SMALL_FIVE_PT)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def classify_heading(text: str) -> str | None:
    normalized = normalize_text(text)
    if FIRST_LEVEL_HEADING_RE.match(normalized):
        return "heading1"
    if SECOND_LEVEL_HEADING_RE.match(normalized):
        return "heading2"
    if THIRD_LEVEL_HEADING_RE.match(normalized):
        return "heading3"
    if FOURTH_LEVEL_HEADING_RE.match(normalized) or FIFTH_LEVEL_HEADING_RE.match(normalized):
        return "run_in_heading"
    if TABLE_CAPTION_RE.match(normalized) or FIGURE_CAPTION_RE.match(normalized):
        return "caption"
    if SOURCE_NOTE_RE.match(normalized):
        return "source_note"
    return None


def standardize_document(
    src_path: Path,
    *,
    output_dir: Path | None,
    suffix: str,
    add_header_footer: bool,
    overwrite: bool,
) -> Path:
    document = Document(src_path)
    configure_page(document)
    configure_normal_style(document)

    paragraphs = non_empty_paragraphs(document)
    if not paragraphs:
        raise ValueError(f"{src_path.name} 没有可处理的正文内容。")

    title_index = 0
    english_title_index = 1 if len(paragraphs) > 1 and looks_like_english_title(paragraphs[1].text) else None
    abstract_index = find_paragraph_index(paragraphs, lambda p: is_abstract_heading(p.text))
    english_abstract_index = find_paragraph_index(paragraphs, lambda p: is_english_abstract_heading(p.text))
    keywords_index = find_paragraph_index(paragraphs, lambda p: is_keywords_line(p.text))
    english_keywords_index = find_paragraph_index(paragraphs, lambda p: is_english_keywords_line(p.text))
    references_index = find_paragraph_index(paragraphs, lambda p: is_reference_heading(p.text))

    abstract_body_end = next_index([keywords_index, english_abstract_index, references_index, len(paragraphs)])
    english_abstract_body_end = next_index([english_keywords_index, references_index, len(paragraphs)])

    for index, paragraph in enumerate(paragraphs):
        text = normalize_text(paragraph.text)
        has_math = paragraph_has_math(paragraph)
        if not text and not has_math:
            continue

        if index == title_index:
            apply_title_format(paragraph)
            continue
        if english_title_index is not None and index == english_title_index:
            apply_english_title_format(paragraph)
            continue
        if abstract_index is not None and index == abstract_index:
            apply_abstract_heading_format(paragraph)
            continue
        if english_abstract_index is not None and index == english_abstract_index:
            apply_english_abstract_heading_format(paragraph)
            continue
        if keywords_index is not None and index == keywords_index:
            apply_keywords_format(paragraph, english=False)
            continue
        if english_keywords_index is not None and index == english_keywords_index:
            apply_keywords_format(paragraph, english=True)
            continue
        if references_index is not None and index == references_index:
            apply_reference_heading_format(paragraph)
            continue

        if abstract_index is not None and abstract_body_end is not None and abstract_index < index < abstract_body_end:
            apply_abstract_body_format(paragraph, english=False)
            continue
        if (
            english_abstract_index is not None
            and english_abstract_body_end is not None
            and english_abstract_index < index < english_abstract_body_end
        ):
            apply_abstract_body_format(paragraph, english=True)
            continue
        if references_index is not None and index > references_index:
            apply_reference_item_format(paragraph)
            continue
        if has_math and not text:
            apply_equation_block_format(paragraph)
            continue

        heading_type = classify_heading(text)
        if heading_type == "heading1":
            apply_level1_heading_format(paragraph)
        elif heading_type == "heading2":
            apply_level2_heading_format(paragraph)
        elif heading_type == "heading3":
            apply_level3_heading_format(paragraph)
        elif heading_type == "run_in_heading":
            apply_run_in_heading_format(paragraph)
        elif heading_type == "caption":
            apply_caption_format(paragraph)
        elif heading_type == "source_note":
            apply_source_note_format(paragraph)
        else:
            apply_body_format(paragraph)

    format_table_cells(document)
    set_header_footer(document, normalize_text(paragraphs[title_index].text), enabled=add_header_footer)

    target_dir = output_dir if output_dir is not None else src_path.parent
    target_dir.mkdir(parents=True, exist_ok=True)
    output_path = target_dir / f"{src_path.stem}{suffix}{src_path.suffix}"
    if output_path.exists() and not overwrite:
        raise FileExistsError(f"输出文件已存在，请使用 --overwrite 或调整后缀: {output_path}")
    document.save(output_path)
    return output_path


def choose_option(cli_value, config_value):
    return config_value if cli_value is None else cli_value


def run_batch(
    paths: Sequence[str | Path],
    *,
    output_dir: Path | None = None,
    suffix: str | None = None,
    overwrite: bool | None = None,
    add_header_footer: bool | None = None,
    config_path: Path | None = None,
) -> tuple[list[Path], list[str]]:
    _, config = load_config(config_path)
    apply_runtime_config(config)

    resolved_suffix = suffix or config["general"]["suffix"]
    resolved_overwrite = choose_option(overwrite, config["general"]["overwrite"])
    resolved_add_header_footer = choose_option(add_header_footer, config["general"]["add_header_footer"])

    if not paths:
        return [], ["未提供要格式化的文档路径。"]

    input_files = list(iter_input_files(paths, resolved_suffix))
    if not input_files:
        return [], ["没有找到可处理的 DOCX 文件。"]

    resolved_output_dir = resolve_output_dir(input_files, output_dir)
    produced: list[Path] = []
    errors: list[str] = []

    for src_path in input_files:
        try:
            output_path = standardize_document(
                src_path,
                output_dir=resolved_output_dir,
                suffix=resolved_suffix,
                add_header_footer=resolved_add_header_footer,
                overwrite=resolved_overwrite,
            )
            produced.append(output_path)
        except Exception as exc:  # pragma: no cover - batch mode reports item-level failures
            errors.append(f"处理失败 {src_path}: {exc}")

    return produced, errors
